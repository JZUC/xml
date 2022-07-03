import docx
import pandas as pd
from set_cell_border import set_cell_border
from list_number import list_number
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE

def change_orientation():
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section

# Create the dataframe the csv file (C://test.cv) containing the XML data and
dfAlarms = pd.read_csv('C:/Users/jzuccaro/Desktop/test.csv', encoding="ISO-8859-1")

# Create Document object
doc = docx.Document()

#### HEADER ####

section = doc.sections[0]
header = section.header
hp = header.paragraphs[0]
hp.text = "Left Text\tCenter Text\tRight Text"
hp.style = doc.styles["Header"]
hp.paragraph_format.left_indent = -Inches(1.0)
hp.paragraph_format.right_indent = Inches(0.5)

### Continue research here:
### https://python-docx.readthedocs.io/en/latest/user/documents.html


#######


#### STYLES ####

# Create styles object and add a new style to the document's options
styles = doc.styles
style = styles.add_style('test', WD_STYLE_TYPE.PARAGRAPH)

#Set text font and English font: Times New Roman
doc.styles['test'].font.name = 'Georgia'

# Set custom style text font size
doc.styles['test'].font.size = Pt(9)

# Set custom text font color
doc.styles['test'].font.color.rgb = RGBColor(0, 0, 0)


test = doc.add_paragraph(text ="Georgia 9", style ="test")

paragraphs = doc.paragraphs
paragraph = paragraphs[0]
paragraph.style = doc.styles["test"]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

########



#### BULLET LIST ####

p0 = doc.add_paragraph('')
run = p0.add_run('Level 1, Georgia 9')
run.font.name = 'Georgia'
run.font.size = Pt(9)
p0.style = 'List Bullet'

p1 = doc.add_paragraph('')
run = p1.add_run('Level 2, Georgia 9')
run.font.name = 'Georgia'
run.font.size = Pt(9)
p1.style = 'List Bullet 2'

p2 = doc.add_paragraph('')
run = p2.add_run('Level 3, Georgia 9')
run.font.name = 'Georgia'
run.font.size = Pt(9)
p2.style = 'List Bullet 3'

#### NUMBER LIST ####

p0 = doc.add_paragraph('')
run = p0.add_run('Level 1, Georgia 9')
run.font.name = 'Georgia'
run.font.size = Pt(9)
p0.style = 'List Number'


p1 = doc.add_paragraph('')
run = p1.add_run('Level 2, Georgia 9')
run.font.name = 'Georgia'
run.font.size = Pt(9)
p1.style = 'List Number 2'

p2 = doc.add_paragraph('')
run = p2.add_run('Level 3, Georgia 9')
run.font.name = 'Georgia'
run.font.size = Pt(9)
p2.style = 'List Number 3'

########



#### ALARM TABLE ####

change_orientation()

# Add a table to the doc object. Its dimensions will be equal to dfAlarms with an extra row added for inclusion of
# a header.
# dfAlarms.shape[0] represents the dataframe's number of rows while shape[1] represents columns.
almTable = doc.add_table(dfAlarms.shape[0]+1, dfAlarms.shape[1])

# Add column headers to almTable
for i in range(dfAlarms.shape[1]):
    almTable.cell(0, i).text = dfAlarms.columns[i]

# Fill in the rest of almTable using values from dfAlarms
for i in range(dfAlarms.shape[0]):
    for j in range(dfAlarms.shape[1]):
        almTable.cell(i+1, j).text = str(dfAlarms.values[i, j])

# Apply Table Grid style to almTable
## almTable.style = "Table Grid"

# Create hdr object containing column titles
hdr = almTable.rows[0].cells

# Reformat font across entire table
for row in almTable.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size = docx.shared.Pt(8)



#### TABLE HEADER ####

# Change the font of the top row of almTable
# i represents the column number
for i in range(dfAlarms.shape[1]):
    paragraph = hdr[i].paragraphs[0]
    run = paragraph.runs
    print(paragraph)
    font = run[0].font
    font.size = docx.shared.Pt(10)  # font size = 30
    font.bold = True

rows = dfAlarms.shape[0]
cols = dfAlarms.shape[1]
n1 = cols-2

i = 0

while i < n1:
    set_cell_border(
            almTable.cell(0, i),
            top = {"sz": 15, "val": "thick", "color": "#000000", "space": "0"},
            bottom = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"}
            )
    i += 1

i = cols - 2

while i < cols:
    set_cell_border(
            almTable.cell(0, i),
            top = {"sz": 15, "val": "thick", "color": "#000000", "space": "0"},
            bottom = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"},
            start = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"}
            )
    i += 1

i = 0

#### END HEADER ####

# Format cell borders of last two columns
while i < rows:
    set_cell_border(
            almTable.cell(i, cols-2),
            bottom = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"},
            start = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"}
            )

    set_cell_border(
            almTable.cell(i, cols-1),
            bottom = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"},
            start = {"sz": 10, "val": "thick", "color": "#000000", "space": "0"}
            )
    i += 1


# save the doc
doc.save('./test.docx')